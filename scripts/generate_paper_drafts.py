from __future__ import annotations

import sys
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from zipfile import ZipFile

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "papers"
OUT_DIR.mkdir(parents=True, exist_ok=True)

EN_OUT = OUT_DIR / "MPI_Science_Style_Draft_EN.docx"
ZH_OUT = OUT_DIR / "MPI_煤炭学报格式_中文初稿.docx"

FIGURES = [
    {
        "path": ROOT / "frontend" / "public" / "mpi-algorithm" / "flow_overview.png",
        "en": "Figure 1. Workflow of the new MPI research pipeline.",
        "zh": "图1 新MPI科研流程总览。",
    },
    {
        "path": ROOT / "frontend" / "public" / "mpi-algorithm" / "rsi_stability.png",
        "en": "Figure 2. RSI-PhaseField stability behavior.",
        "zh": "图2 RSI相场稳定性响应。",
    },
    {
        "path": ROOT / "frontend" / "public" / "mpi-algorithm" / "bri_depth_curve.png",
        "en": "Figure 3. BRI trend under depth and energy factors.",
        "zh": "图3 BRI随埋深与能量因子变化趋势。",
    },
    {
        "path": ROOT / "frontend" / "public" / "mpi-algorithm" / "asi_stress_profile.png",
        "en": "Figure 4. ASI-UST stress profile.",
        "zh": "图4 ASI-UST应力剖面。",
    },
    {
        "path": ROOT / "frontend" / "public" / "mpi-algorithm" / "mpi_panels.png",
        "en": "Figure 5. RSI/BRI/ASI/MPI panels.",
        "zh": "图5 RSI/BRI/ASI/MPI综合面板。",
    },
]


def _set_style_font(style: Any, ascii_font: str, east_asia_font: str, size_pt: int, bold: bool = False) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size_pt)
    style.font.bold = bold
    style._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)



def _setup_page(doc: Document, a4: bool = True) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        if a4:
            section.page_width = Cm(21.0)
            section.page_height = Cm(29.7)



def _configure_en_styles(doc: Document) -> None:
    _set_style_font(doc.styles["Normal"], "Times New Roman", "Times New Roman", 12)
    _set_style_font(doc.styles["Heading 1"], "Times New Roman", "Times New Roman", 14, bold=True)
    _set_style_font(doc.styles["Heading 2"], "Times New Roman", "Times New Roman", 12, bold=True)
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        p = doc.styles[style_name].paragraph_format
        p.line_spacing = 2.0
        p.space_before = Pt(0)
        p.space_after = Pt(6)



def _configure_zh_styles(doc: Document) -> None:
    _set_style_font(doc.styles["Normal"], "Times New Roman", "宋体", 12)
    _set_style_font(doc.styles["Heading 1"], "Times New Roman", "黑体", 14, bold=True)
    _set_style_font(doc.styles["Heading 2"], "Times New Roman", "黑体", 12, bold=True)
    for style_name in ("Normal", "Heading 1", "Heading 2"):
        p = doc.styles[style_name].paragraph_format
        p.line_spacing = 1.5
        p.space_before = Pt(0)
        p.space_after = Pt(6)



def _add_center_text(doc: Document, text: str, size: int = 12, bold: bool = False, east_asia_font: str = "宋体") -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia_font)



def _add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Cm(15.5))
    cp = doc.add_paragraph(caption)
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER



def _add_omml_equation(doc: Document, equation_text: str, number: Optional[str] = None) -> None:
    """
    Insert a Word OMML equation object instead of plain pseudo-formula text.
    """
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    omath_para = OxmlElement("m:oMathPara")
    omath = OxmlElement("m:oMath")
    mr = OxmlElement("m:r")
    mt = OxmlElement("m:t")
    mt.text = equation_text
    mr.append(mt)
    omath.append(mr)
    omath_para.append(omath)
    p._p.append(omath_para)

    if number:
        n = doc.add_paragraph(number)
        n.alignment = WD_ALIGN_PARAGRAPH.RIGHT


def _build_demo_point(name: str, depth: float, coal_thickness: float, roof_e: float, roof_tensile: float, friction: float) -> Dict[str, Any]:
    return {
        "name": name,
        "depth": depth,
        "point": {
            "x": 0.0,
            "y": 0.0,
            "borehole": name,
            "thickness": coal_thickness,
            "burial_depth": depth,
            "strata": [
                {
                    "name": "sandstone_roof",
                    "thickness": 7.0,
                    "elastic_modulus": roof_e,
                    "tensile_strength": roof_tensile,
                    "cohesion": 4.0,
                    "friction_angle": friction,
                    "compressive_strength": 65.0,
                },
                {
                    "name": "mudstone_interlayer",
                    "thickness": 4.0,
                    "elastic_modulus": 9.0,
                    "tensile_strength": 1.3,
                    "cohesion": 2.0,
                    "friction_angle": 27.0,
                    "compressive_strength": 30.0,
                },
            ],
        },
    }



def _run_new_algorithm_demo() -> Optional[Dict[str, Any]]:
    try:
        if str(ROOT / "backend") not in sys.path:
            sys.path.insert(0, str(ROOT / "backend"))
        if str(ROOT) not in sys.path:
            sys.path.insert(0, str(ROOT))

        from app.services.mpi_calculator import MPIConfig, PointData, RockLayer, calc_all_indicators
        from mpi_advanced.core.data_models import GeologyLayer, GeologyLayerType, GeologyModel, MiningParameters
        from mpi_advanced.indicators.asi_indicator_ust import ASIIndicatorUST

        cfg = MPIConfig(ust_parameter_b=0.5, phase_field_length_scale=0.5)
        scenarios = [
            _build_demo_point("Scenario-LowRisk", 380.0, 3.2, 24.0, 3.8, 34.0),
            _build_demo_point("Scenario-MidRisk", 460.0, 3.8, 18.0, 2.8, 30.0),
            _build_demo_point("Scenario-HighRisk", 580.0, 4.6, 12.0, 1.8, 26.0),
        ]

        rows: List[Tuple[str, Dict[str, float]]] = []
        for item in scenarios:
            p = item["point"]
            point = PointData(
                x=float(p["x"]),
                y=float(p["y"]),
                borehole=str(p["borehole"]),
                thickness=float(p["thickness"]),
                burial_depth=float(p["burial_depth"]),
                strata=[RockLayer(**layer) for layer in p["strata"]],
            )
            rows.append((item["name"], calc_all_indicators(point, config=cfg)))

        def _geo(roof_thickness: float, support_pressure: float) -> GeologyModel:
            layers = [
                GeologyLayer(
                    name="Coal",
                    layer_type=GeologyLayerType.COAL,
                    thickness=3.5,
                    depth_top=0.0,
                    depth_bottom=3.5,
                    elastic_modulus=3e9,
                    poisson_ratio=0.30,
                    cohesion=2e6,
                    friction_angle=20.0,
                    tensile_strength=0.8e6,
                    fracture_toughness=0.5,
                    density=1400,
                ),
                GeologyLayer(
                    name="Roof",
                    layer_type=GeologyLayerType.SANDSTONE,
                    thickness=roof_thickness,
                    depth_top=3.5,
                    depth_bottom=3.5 + roof_thickness,
                    elastic_modulus=20e9,
                    poisson_ratio=0.25,
                    cohesion=6e6,
                    friction_angle=32.0,
                    tensile_strength=2.5e6,
                    fracture_toughness=1.2,
                    density=2500,
                ),
            ]
            mining = MiningParameters(
                panel_length=180.0,
                panel_width=140.0,
                mining_height=3.5,
                mining_depth=450.0,
                advance_rate=4.0,
                support_pressure=support_pressure,
            )
            return GeologyModel(layers=layers, mining_params=mining)

        target = ASIIndicatorUST(b=0.6)
        geos = [_geo(5.0, 0.22e6), _geo(6.0, 0.28e6), _geo(7.2, 0.35e6), _geo(8.5, 0.42e6)]
        samples = [{"geology": g, "target_asi": float(target.compute(g).value)} for g in geos]
        cal = ASIIndicatorUST(b=0.1).calibrate_b_parameter(samples, bootstrap_rounds=60, seed=2026)

        return {
            "rows": rows,
            "calibration": {
                "best_b": float(cal["best_b"]),
                "rmse": float(cal["rmse"]),
                "ci95": [float(cal["ci95"][0]), float(cal["ci95"][1])],
                "sample_count": int(cal["sample_count"]),
            },
            "note": "Deterministic synthetic scenarios generated by the new MPI algorithm stack.",
        }
    except Exception:
        return None



def _add_indicator_table(doc: Document, title: str, demo_data: Optional[Dict[str, Any]], zh: bool = False) -> None:
    doc.add_heading(title, level=2)
    headers = ["Scenario", "RSI", "BRI", "ASI", "MPI"] if not zh else ["场景", "RSI", "BRI", "ASI", "MPI"]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    if demo_data:
        for scenario, values in demo_data["rows"]:
            row = table.add_row().cells
            row[0].text = scenario
            row[1].text = f"{values['rsi']:.2f}"
            row[2].text = f"{values['bri']:.2f}"
            row[3].text = f"{values['asi']:.2f}"
            row[4].text = f"{values['mpi']:.2f}"

        cal = demo_data["calibration"]
        if zh:
            doc.add_paragraph(
                f"探索性结论：UST参数b校准结果：best_b={cal['best_b']:.3f}, RMSE={cal['rmse']:.3f}, "
                f"95%CI=[{cal['ci95'][0]:.3f}, {cal['ci95'][1]:.3f}], 样本数={cal['sample_count']}。"
            )
            doc.add_paragraph(
                "探索性结论：当前演示数据未进行正式显著性检验（p-value），最终实证将报告显著性与效应量，"
                "并给出相对提升（relative improvement）与Cohen指标。"
            )
            doc.add_paragraph("注：上述结果用于验证新算法流程与排版结构，不作为最终投稿结论。")
        else:
            doc.add_paragraph(
                f"[Exploratory] UST b-parameter calibration: best_b={cal['best_b']:.3f}, RMSE={cal['rmse']:.3f}, "
                f"95%CI=[{cal['ci95'][0]:.3f}, {cal['ci95'][1]:.3f}], samples={cal['sample_count']}."
            )
            doc.add_paragraph(
                "[Exploratory] No formal significance test (p-value) is claimed for draft scenarios; final field-label runs "
                "will report significance and effect size, including relative improvement and Cohen's d."
            )
            doc.add_paragraph("Note: This table validates the new algorithm pipeline and manuscript structure only.")
    else:
        doc.add_paragraph(
            "Placeholder table. Rerun with local dependencies to generate new-algorithm scenario values."
            if not zh
            else "结果占位。请在本地依赖完整时重跑脚本以生成新算法演示结果。"
        )



def build_en_doc(demo_data: Optional[Dict[str, Any]]) -> None:
    doc = Document()
    _setup_page(doc, a4=False)
    _configure_en_styles(doc)

    _add_center_text(
        doc,
        "New MPI Algorithm Draft for SCI Submission: RSI-PhaseField and ASI-UST Integration",
        size=16,
        bold=True,
        east_asia_font="Times New Roman",
    )
    _add_center_text(doc, "Author A, Author B, Author C", size=12, east_asia_font="Times New Roman")
    _add_center_text(doc, "Affiliation 1; Affiliation 2", size=11, east_asia_font="Times New Roman")
    _add_center_text(doc, "One-sentence summary: A reproducible MPI pipeline built on new phase-field and UST algorithms.", size=11, east_asia_font="Times New Roman")

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This manuscript draft is fully based on the new MPI algorithm stack. RSI is computed by a phase-field fracture "
        "formulation with a lightweight 2D finite-difference solver and explicit convergence diagnostics. [Moderate] ASI is computed by "
        "a unified-strength-theory (UST) analytical model with calibratable parameter b and uncertainty interval output. "
        "BRI is handled by a microseismic-ready module that falls back to deterministic physics-aware scoring when no "
        "microseismic stream is available. The implementation is connected to reproducibility endpoints for dataset manifesting, "
        "split auditing, experiment tracking, and artifact archiving. [Exploratory] We provide draft-level scenario results and current figures "
        "to support immediate manuscript writing, while reserving strict real-label field validation for final claims."
    )
    doc.add_paragraph("Keywords: mine pressure; phase-field fracture; unified strength theory; reproducibility; risk assessment")

    doc.add_heading("Main Text", level=1)
    doc.add_heading("1. Introduction", level=2)
    doc.add_paragraph(
        "The previous empirical MPI formulas have been retired in this draft. The present version only uses the new research-oriented "
        "algorithm implementations from the mpi_advanced module."
    )

    doc.add_heading("2. New MPI Algorithm Stack", level=2)
    doc.add_paragraph("RSI module: RSI-PhaseField (2D finite-difference solver, boundary constraints, convergence checks; Eq. (3)-(4)).")
    doc.add_paragraph("ASI module: ASI-UST (analytical stress field, parameter b calibration with bootstrap CI; Eq. (2), (5)).")
    doc.add_paragraph("BRI module: microseismic-driven architecture with no-signal fallback branch.")
    doc.add_paragraph("Fusion module: weighted MPI synthesis under explicit reproducibility metadata (Eq. (1)).")
    _add_omml_equation(
        doc,
        "MPI = w_RSI * RSI + w_BRI * BRI + w_ASI * ASI",
        number="(1)",
    )
    _add_omml_equation(
        doc,
        "F_UST = σ₁ - (σ₂ + b * σ₃)/(1 + b) - f_t",
        number="(2)",
    )
    _add_omml_equation(
        doc,
        "- l₀² * ∇²ϕ + ϕ = 1 - α * S(x,y)",
        number="(3)",
    )
    _add_omml_equation(
        doc,
        "D = 1 - ϕ",
        number="(4)",
    )
    _add_omml_equation(
        doc,
        "CI95(b) = [Q_2.5%(b*), Q_97.5%(b*)]",
        number="(5)",
    )

    doc.add_heading("3. Reproducibility and Evidence Protocol", level=2)
    doc.add_paragraph(
        "All experiments are designed to bind dataset version, split manifest, seed, and model spec. Metrics and artifacts are archived "
        "for traceability and manuscript reconstruction."
    )
    doc.add_paragraph(
        "[Moderate] Statistical reporting protocol includes uncertainty, significance (p-value), and effect size requirements "
        "for submission-stage experiments."
    )
    doc.add_paragraph("Calibration uncertainty is reported by Eq. (5), and all risk fusion outputs are tied to Eq. (1).")

    _add_indicator_table(doc, "4. New-Algorithm Demonstration Results", demo_data, zh=False)

    doc.add_heading("5. Current Figures", level=2)
    for fig in FIGURES:
        _add_figure(doc, fig["path"], fig["en"])

    doc.add_heading("6. Limitations and Next Submission Step", level=2)
    doc.add_paragraph(
        "Current values are scenario demonstrations for pipeline verification. Final paper claims must be based on strict field-label "
        "experiments using frozen dataset manifests and leakage-safe splits."
    )

    doc.add_heading("References", level=1)
    refs = [
        "Francfort, G. A., & Marigo, J.-J. (1998). Revisiting brittle fracture as an energy minimization problem. DOI:10.1016/S0022-5096(98)00034-9.",
        "Miehe, C., Hofacker, M., & Welschinger, F. (2010). A phase field model for rate-independent crack propagation. DOI:10.1016/j.cma.2010.04.011.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    doc.save(str(EN_OUT))



def build_zh_doc(demo_data: Optional[Dict[str, Any]]) -> None:
    doc = Document()
    _setup_page(doc, a4=True)
    _configure_zh_styles(doc)

    _add_center_text(doc, "基于新MPI算法的论文初稿（RSI相场 + ASI-UST）", size=18, bold=True, east_asia_font="黑体")
    _add_center_text(doc, "作者甲，作者乙，作者丙", size=12, east_asia_font="宋体")
    _add_center_text(doc, "（单位名称，城市 邮编）", size=11, east_asia_font="宋体")

    doc.add_paragraph("摘  要：")
    doc.add_paragraph(
        "本文初稿完全基于新MPI算法体系，不再使用旧版经验公式。RSI采用相场断裂模型与二维有限差分求解，"
        "显式输出收敛信息；ASI采用统一强度理论（UST）解析模型，并支持参数b校准与置信区间估计；BRI采用微震驱动架构，"
        "在无微震流条件下自动降级到物理约束分支。系统与研究接口联动，实现数据版本固定、切分审计、实验归档与结果追溯。"
        "探索性结论：文中给出基于新算法的演示性场景结果和当前图件，用于支撑中英文论文排版与结构成稿。"
    )
    doc.add_paragraph("关键词：矿压影响指标；相场断裂；统一强度理论；可复现性；风险评估")

    doc.add_heading("1 引言", level=1)
    doc.add_paragraph("旧版MPI经验公式已在本轮中退出主流程，当前后端计算与论文材料统一切换到新MPI算法实现。")

    doc.add_heading("2 新MPI算法体系", level=1)
    doc.add_heading("2.1 RSI相场模块", level=2)
    doc.add_paragraph("采用二维有限差分相场求解器，包含边界条件、迭代收敛判据与损伤指标提取，见式（3）-（4）。")

    doc.add_heading("2.2 ASI-UST模块", level=2)
    doc.add_paragraph("采用统一强度理论解析解，支持参数b网格搜索与Bootstrap置信区间，见式（2）与式（5）。")

    doc.add_heading("2.3 BRI模块", level=2)
    doc.add_paragraph("采用微震驱动架构，在缺失微震流时走受约束降级分支，保证工程可用性与模型一致性。")
    _add_omml_equation(
        doc,
        "MPI = w_RSI * RSI + w_BRI * BRI + w_ASI * ASI",
        number="（1）",
    )
    _add_omml_equation(
        doc,
        "F_UST = σ₁ - (σ₂ + b * σ₃)/(1 + b) - f_t",
        number="（2）",
    )
    _add_omml_equation(
        doc,
        "- l₀² * ∇²ϕ + ϕ = 1 - α * S(x,y)",
        number="（3）",
    )
    _add_omml_equation(
        doc,
        "D = 1 - ϕ",
        number="（4）",
    )
    _add_omml_equation(
        doc,
        "CI95(b) = [Q_2.5%(b*), Q_97.5%(b*)]",
        number="（5）",
    )

    doc.add_heading("3 可复现性协议", level=1)
    doc.add_paragraph("中等结论：通过 /api/research/* 接口完成数据注册、切分、实验运行与产物追踪，保证图表和结论可追溯。")
    doc.add_paragraph(
        "中等结论：统计报告将覆盖显著性（p-value）与效应量，并在正式实验中给出相对提升（relative improvement）。"
    )

    _add_indicator_table(doc, "4 新算法演示结果", demo_data, zh=True)

    doc.add_heading("5 当前图件", level=1)
    for fig in FIGURES:
        _add_figure(doc, fig["path"], fig["zh"])

    doc.add_heading("6 讨论与后续", level=1)
    doc.add_paragraph(
        "当前结果用于验证新算法流程与论文结构，最终投稿结论需基于真实矿压标签与严格切分实验。"
    )

    doc.add_heading("参考文献", level=1)
    refs = [
        "Francfort G A, Marigo J J. Revisiting brittle fracture as an energy minimization problem. DOI:10.1016/S0022-5096(98)00034-9.",
        "Miehe C, Hofacker M, Welschinger F. A phase field model for rate-independent crack propagation. DOI:10.1016/j.cma.2010.04.011.",
    ]
    for i, ref in enumerate(refs, 1):
        doc.add_paragraph(f"{i}. {ref}")

    doc.save(str(ZH_OUT))



def _count_omml_equations(docx_path: Path) -> int:
    with ZipFile(docx_path, "r") as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    return xml.count("<m:oMath")


def _extract_document_xml(docx_path: Path) -> str:
    with ZipFile(docx_path, "r") as zf:
        return zf.read("word/document.xml").decode("utf-8", errors="ignore")


def _verify_equation_numbering(xml: str, labels: List[str]) -> bool:
    return all(label in xml for label in labels)


def main() -> None:
    demo_data = _run_new_algorithm_demo()
    build_en_doc(demo_data)
    build_zh_doc(demo_data)
    en_eq = _count_omml_equations(EN_OUT)
    zh_eq = _count_omml_equations(ZH_OUT)
    en_xml = _extract_document_xml(EN_OUT)
    zh_xml = _extract_document_xml(ZH_OUT)
    en_labels_ok = _verify_equation_numbering(en_xml, ["(1)", "(2)", "(3)", "(4)", "(5)"])
    lpar = chr(0xFF08)
    rpar = chr(0xFF09)
    zh_labels_ok = _verify_equation_numbering(
        zh_xml,
        [f"{lpar}1{rpar}", f"{lpar}2{rpar}", f"{lpar}3{rpar}", f"{lpar}4{rpar}", f"{lpar}5{rpar}"],
    )
    if en_eq < 5 or zh_eq < 5 or not en_labels_ok or not zh_labels_ok:
        raise RuntimeError(f"Equation QA failed: EN={en_eq}, ZH={zh_eq}")
    print(f"Generated: {EN_OUT}")
    print(f"Generated: {ZH_OUT}")


if __name__ == "__main__":
    main()
